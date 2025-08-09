# docx_utils.py
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from typing import List, Tuple

def extract_paragraphs_from_docx(path: str) -> List[str]:
    doc = Document(path)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs

def extract_structured_text(path: str) -> List[Tuple[int, str]]:
    """
    Returns list of (paragraph_index, text)
    """
    doc = Document(path)
    out = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            out.append((i, text))
    return out

def insert_comment_simulation(in_path: str, out_path: str, annotations: list):
    """
    annotations: list of dicts:
      {
        "paragraph_index": int,
        "match_text": str or None,
        "comment": str
      }
    We will:
      - load doc
      - for target paragraph index: find the run containing match_text if provided (fallback highlight entire paragraph)
      - highlight that run (by setting shading/background via run._r element or add appended comment text)
      - append an inline bracketed comment at end of paragraph
    """
    doc = Document(in_path)
    for ann in annotations:
        idx = ann.get("paragraph_index")
        comment = ann.get("comment", "")
        match_text = ann.get("match_text")

        if idx < 0 or idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]

        # try to find run containing match_text
        target_run = None
        if match_text:
            for run in para.runs:
                if match_text in run.text:
                    target_run = run
                    break

        if target_run is None:
            # fallback: highlight first run or whole paragraph by creating a new run
            target_run = para.add_run("")  # empty run, we'll add highlight by adding comment text

        # append bracketed comment
        # Add a new run for comment
        comment_run = para.add_run(f"  [COMMENT: {comment}]")
        font = comment_run.font
        font.italic = True
        font.size = Pt(9)

        # We can also mark the target run by surrounding with markers (for visibility)
        # If match_text present, wrap it with markers
        if match_text:
            target_run.text = target_run.text  # leave text
        else:
            # add a small visible marker at start
            para.insert_paragraph_before("")  # ensure no None
    doc.save(out_path)


def simple_highlight_paragraph(in_path: str, out_path: str, paragraph_indices: list):
    """
    Create a copy where paragraphs listed are followed by a [FLAGGED] note.
    """
    doc = Document(in_path)
    for idx in paragraph_indices:
        if idx < 0 or idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        para.add_run(f"  [FLAGGED]").italic = True
    doc.save(out_path)


def sanitize_filename(name: str) -> str:
    return re.sub(r"[^\w\-_\. ]", "_", name)
